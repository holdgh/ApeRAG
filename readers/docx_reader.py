import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, cast

from docx import Document as Docx_document
from llama_index.readers.base import BaseReader
from llama_index.schema import Document

CHUNK_SPLIT_THRESHOLD = 500


class Stacks:
    """
    An array of stacks.
    the array index is the docx title level
    the every stack store the level contents
    """

    def __init__(self):
        self.stacks = [[]]  # [] is a placeholder

    def push(self, level: int, value: str):
        while level >= len(self.stacks):
            self.stacks.append([])
        self.stacks[level].append(value)

    def pop(self, level: int):
        if level >= len(self.stacks):
            return None
        if len(self.stacks[level]) == 0:
            return None
        return self.stacks[level].pop()

    def package_content(self, level: int):
        """
        package the stack contents to a trunk from  0 ~ level
        :param level: the deepest level to package
        :return: content string
        """
        res = ""
        for i in range(0, level + 1):
            for j in range(0, len(self.stacks[i])):
                res += self.stacks[i][j]
        return res

    def count_contents(self, level: int):
        """
        count the contexts for level
        :param level: the stack level we want to count
        :return: the total length of content at level
        """
        res = 0
        for i in range(0, len(self.stacks[level])):
            res += len(self.stacks[level][i])
        return res

    def remove(self, level):
        """
        remove the stacks contents from level to the deepest level
        :param level: begin level
        """
        for i in range(level, len(self.stacks)):
            while len(self.stacks[i]) > 0:
                self.pop(i)

    def get_title(self, level):
        """
        get the title
        :param level: the level of the title
        :return: the content of title
        """
        return self.stacks[level][0]


class MyDocxReader(BaseReader):
    """Docx parser."""

    def __init__(self):
        super().__init__()
        self.contents_stacks = Stacks()

    def phrase_level(self, style="Normal"):
        if style == 'Normal':  # content level set 0
            return 0
        level = re.search(r'\d+', style)  # title level set 1:
        if level:
            return int(level.group())
        return 0

    def docx_to_tups(self, docx_path: Path) -> List[Tuple[Optional[List[str]], str]]:
        docx_tups: List[Tuple[Optional[List[str]], str]] = []  # [(title,content)]
        obj = Docx_document(docx_path)
        current_level = 0  # 当前级别
        headers = 0  # 专门记录当前级别title的个数
        for p in obj.paragraphs:
            style = p.style.name
            level = self.phrase_level(style)
            if 0 < level <= 5:  # docx 标题支持自定义级别 只关注1-5级标题
                if level < current_level:  # 遇到一个新的父标题 打包之前的子标题内容
                    context = self.contents_stacks.package_content(current_level)
                    docx_tups.append((self.contents_stacks.get_title(level), context))
                    self.contents_stacks.remove(level)  # 将level级别到current_level级别的内容出栈
                    headers = 1  # 重置

                elif level == current_level:  # 遇到同级标题 判断该同级别的内容长度是否满足 trunk
                    if self.contents_stacks.count_contents(level) >= CHUNK_SPLIT_THRESHOLD:  # 满足 组成新的chunk
                        if headers > 1:
                            title = self.contents_stacks.get_title(level - 1)  # 多个同级标题组成chunk 取其父标题
                        else:
                            title = self.contents_stacks.get_title(level)
                        context = self.contents_stacks.package_content(current_level)
                        docx_tups.append((title, context))
                        self.contents_stacks.remove(level)
                        self.contents_stacks.push(level, p.text)  # 新内容入栈
                        headers = 1
                    else:
                        self.contents_stacks.push(current_level, "\n" + p.text)  # 不满足 继续添加
                        headers += 1
                else:  # 遇到子标题 继续遍历更深的内容
                    headers = 1
                    current_level = level
                    self.contents_stacks.push(current_level, p.text)
            else:  # 遇到正文 加入当前的level 内容栈中
                self.contents_stacks.push(current_level, "\n" + p.text)

        if headers > 1 and current_level > 1:
            title = self.contents_stacks.get_title(current_level - 1)  # 多个同级标题组成chunk 取其父标题
        else:
            title = self.contents_stacks.get_title(current_level)
        docx_tups.append((title, context))  # the last chunk
        return docx_tups

    def load_data(self, file: Path, metadata: Optional[Dict] = None) -> List[Document]:
        # """Parse file."""
        try:
            import docx2txt
        except ImportError:
            raise ImportError(
                "docx2txt is required to read Microsoft Word files: "
                "`pip install docx2txt`"
            )

        text = docx2txt.process(file)
        metadata = {"file_name": file.name}
        if len(text) > 5000:  # string length, not world counts
            if metadata is not None:
                metadata.update(metadata)
            return [Document(text=text, metadata=metadata or {})]

        tups = self.docx_to_tups(file)
        results = []
        for header, value in tups:
            text = f"\n{header}\n{value}"
            results.append(Document(text=text, metadata=metadata or {}))
        return results
