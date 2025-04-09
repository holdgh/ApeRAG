import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import docx
from docx.document import Document as Docx_document
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from llama_index.core.readers.base import BaseReader
from llama_index.core.schema import Document

from aperag.utils.utils import Stacks

CHUNK_SPLIT_THRESHOLD = 500


class MyDocxReader(BaseReader):
    """Docx parser."""

    def __init__(self):
        super().__init__()
        self.contents_stacks = Stacks()
        self.current_level = 0  # 当前级别
        self.headers = 0  # 专门记录当前级别title的个数
        self.docx_tups: List[Tuple[Optional[List[str]], str]] = []  # [(title,content)]

    def phrase_level(self, style="Normal"):
        if style == 'Normal' or style == 'Body Text':  # content level set 0  may need fit more format
            return 0
        level = re.search(r'\d+', style)  # title level set 1:
        if level:
            return int(level.group())
        return -1  # 其他类型 如题注（表1. 图1.等）

    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """
        if isinstance(parent, Docx_document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                table = Table(child, parent)
                yield table

    def package_paragraphs(self, paragraph):
        style = paragraph.style.name
        level = self.phrase_level(style)
        if 0 < level <= 5:  # docx 标题支持自定义级别 只关注1-5级标题
            if level < self.current_level:  # 遇到一个新的父标题 打包之前的子标题内容
                context = self.contents_stacks.package_content(self.current_level)
                self.docx_tups.append((self.contents_stacks.get_title(level), context))
                self.contents_stacks.remove(level)  # 将level级别到current_level级别的内容出栈
                self.headers = 1  # 重置

            elif level == self.current_level:  # 遇到同级标题 判断该同级别的内容长度是否满足 trunk
                if self.contents_stacks.count_contents(level) >= CHUNK_SPLIT_THRESHOLD:  # 满足 组成新的chunk
                    if self.headers > 1:
                        title = self.contents_stacks.get_title(level - 1)  # 多个同级标题组成chunk 取其父标题
                    else:
                        title = self.contents_stacks.get_title(level)
                    context = self.contents_stacks.package_content(self.current_level)
                    self.docx_tups.append((title, context))
                    self.contents_stacks.remove(level)
                    self.contents_stacks.push(level, paragraph.text)  # 新内容入栈
                    self.headers = 1
                else:
                    self.contents_stacks.push(self.current_level, "\n" + paragraph.text)  # 不满足 继续添加
                    self.headers += 1
            else:  # 遇到子标题 继续遍历更深的内容
                self.headers = 1
                self.current_level = level
                self.contents_stacks.push(self.current_level, paragraph.text)
        elif level == 0:  # 遇到正文 加入当前的level 内容栈中
            self.contents_stacks.push(self.current_level, "\n" + paragraph.text)
        else:
            return paragraph.text  # 可能是题注 如表名或图名
        return ""  # 如果是正文 则不会计入table的信息

    def package_tables(self, table, last_p, next_p):
        """
        对table类型对象进行解析，last_p是文章中相对table出现的上一自然段，next_p为文章中相对table出现的下一自然段
        :param table: table对象
        :param last_p: 一个string类型的对象
        :param next_p: 一个document child类型的对象，若不存在为None
        :return:
        """
        title = ""
        if next_p is None:
            title = last_p
        elif isinstance(next_p, Paragraph) and self.phrase_level(next_p.style.name) == -1:
            title = next_p.text
        table_head = []
        contexts = ""
        for row in range(0, len(table.rows)):  # 遍历每一行
            row_contents = []
            for col in range(0, len(table.row_cells(row))):  # 遍历行的每一列
                if row == 0:  # 记录表头信息
                    table_head.append(table.cell(row, col).text.replace("\n", ' '))
                else:
                    row_contents.append(table_head[col] + ":" + table.cell(row, col).text.replace("\n", ' '))

            res = " ".join(row_contents)
            contexts += "\n" + res
        self.docx_tups.append((title, contexts))

    def docx_to_tups(self, docx_path: Path) -> List[Tuple[Optional[List[str]], str]]:
        obj = docx.Document(docx_path)
        last_p = ""
        document = self.iter_block_items(obj)
        for block in document:
            if isinstance(block, Paragraph):
                last_p = self.package_paragraphs(block)

            elif isinstance(block, Table):
                next_p = next(document, None)
                self.package_tables(block, last_p, next_p)

        # 处理最后一个chunks
        if self.headers > 1 and self.current_level > 1:
            title = self.contents_stacks.get_title(self.current_level - 1)  # 多个同级标题组成chunk 取其父标题
        else:
            title = self.contents_stacks.get_title(self.current_level)
        self.docx_tups.append((title, self.contents_stacks.package_content(self.current_level)))  # the last chunk

        return self.docx_tups

    def load_data(self, file: Path, metadata: Optional[Dict] = None) -> List[Document]:
        # """Parse file."""
        try:
            import docx2txt
        except ImportError:
            raise ImportError(
                "docx2txt is required to read Microsoft Word files: "
                "`pip install docx2txt`"
            )

        text = docx2txt.process(file)
        metadata = {"file_name": file.name}
        if len(text) < 3000:  # string length, not world counts
            if metadata is not None:
                metadata.update(metadata)
            return [Document(text=text, metadata=metadata or {})]

        tups = self.docx_to_tups(file)
        results = []
        for header, value in tups:
            text = f"\n{header}\n{value}"
            results.append(Document(text=text, metadata=metadata or {}))
        return results
